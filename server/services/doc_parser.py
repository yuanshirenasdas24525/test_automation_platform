"""文档解析服务 — 支持 PDF / Word / Markdown / 纯文本。

职责：
  1. 根据文件扩展名或 MIME 类型分发到对应解析器
  2. 提取纯文本内容（保留段落结构和标题）
  3. 对大文档自动分块（chunk），每块保留上下文元信息

依赖：
  - pypdf       → pip install pypdf --break-system-packages
  - python-docx → pip install python-docx --break-system-packages
  - markdownify → pip install markdownify --break-system-packages (optional)
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 数据结构
# ---------------------------------------------------------------------------


@dataclass
class DocumentChunk:
    """文档的一个逻辑切块。"""
    index: int                       # 块序号（从 0 开始）
    text: str                        # 纯文本内容
    title: str = ""                  # 块所在的小节标题（如有）
    page_range: Tuple[int, int] = (0, 0)  # (起始页, 结束页)


@dataclass
class ParsedDocument:
    """解析后的文档对象。"""
    filename: str
    file_type: str                   # pdf / docx / md / txt
    total_pages: int = 0
    total_chars: int = 0
    plain_text: str = ""             # 完整纯文本
    chunks: List[DocumentChunk] = field(default_factory=list)
    outline: List[str] = field(default_factory=list)  # 标题大纲


# ---------------------------------------------------------------------------
# 分发入口
# ---------------------------------------------------------------------------

MAX_CHUNK_CHARS = 8000   # 每块最大字符数（留给 LLM token 预算）
MAX_TOTAL_CHARS = 200_000  # 单文档最大处理字符数


def parse_document(
    file_path: str,
    chunk: bool = True,
    max_chunk_chars: int = MAX_CHUNK_CHARS,
    max_total_chars: int = MAX_TOTAL_CHARS,
) -> ParsedDocument:
    """解析文档入口。

    Args:
        file_path: 文件绝对路径
        chunk: 是否需要分块（长文档分析用）
        max_chunk_chars: 每块最大字符数
        max_total_chars: 单文档最大处理字符数

    Returns:
        ParsedDocument
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = path.suffix.lower()
    if suffix in (".pdf",):
        doc = _parse_pdf(str(path), chunk, max_chunk_chars, max_total_chars)
    elif suffix in (".docx", ".doc"):
        doc = _parse_docx(str(path), chunk, max_chunk_chars, max_total_chars)
    elif suffix in (".md", ".markdown"):
        doc = _parse_markdown(str(path), chunk, max_chunk_chars, max_total_chars)
    elif suffix in (".txt", ".text", ""):
        doc = _parse_text(str(path), chunk, max_chunk_chars, max_total_chars)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}（支持: PDF / DOCX / MD / TXT）")

    # 长度截断
    if doc.total_chars > max_total_chars:
        doc.plain_text = _truncate_text(doc.plain_text, max_total_chars)
        doc.total_chars = len(doc.plain_text)
        logger.warning(
            "文档 %s 过长（%d chars），已截断至 %d",
            path.name, doc.total_chars, max_total_chars,
        )

    return doc


def parse_text_content(
    text: str,
    source_name: str = "pasted_text",
    chunk: bool = True,
    max_chunk_chars: int = MAX_CHUNK_CHARS,
) -> ParsedDocument:
    """解析用户直接粘贴的文本内容 —— 与文件上传保持统一输出结构。"""
    doc = ParsedDocument(
        filename=source_name,
        file_type="text",
        total_chars=len(text),
        plain_text=text.strip(),
    )
    # 提取标题大纲
    doc.outline = _extract_headings(text)
    if chunk:
        doc.chunks = _chunk_text(text.strip(), max_chunk_chars)
    return doc


# ---------------------------------------------------------------------------
# 各格式解析器
# ---------------------------------------------------------------------------

def _parse_pdf(
    path: str, chunk: bool, max_chunk: int, max_total: int
) -> ParsedDocument:
    """PDF → 纯文本。"""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf 未安装。请执行: pip install pypdf --break-system-packages")

    reader = PdfReader(path)
    pages_text: List[str] = []
    total_pages = len(reader.pages)

    for page_num, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages_text.append(text)
        if sum(len(t) for t in pages_text) > max_total:
            break

    full_text = "\n\n".join(pages_text)
    doc = ParsedDocument(
        filename=Path(path).name,
        file_type="pdf",
        total_pages=total_pages,
        total_chars=len(full_text),
        plain_text=full_text,
    )
    doc.outline = _extract_headings(full_text)
    if chunk:
        doc.chunks = _chunk_text(full_text, max_chunk)

    return doc


def _parse_docx(
    path: str, chunk: bool, max_chunk: int, max_total: int
) -> ParsedDocument:
    """DOCX → 纯文本。保留标题/段落结构。"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx 未安装。请执行: pip install python-docx --break-system-packages"
        )

    docx = Document(path)
    paragraphs: List[str] = []
    for para in docx.paragraphs:
        text = para.text.strip()
        if not text:
            paragraphs.append("")  # 保留空行作为段落分隔
            continue
        # 标题样式 → 加 # 前缀
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style or "title" in style:
            level = 1
            for digit_match in re.findall(r"\d+", style):
                level = int(digit_match)
                break
            paragraphs.append("#" * level + " " + text)
        else:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)
    doc = ParsedDocument(
        filename=Path(path).name,
        file_type="docx",
        total_chars=len(full_text),
        plain_text=full_text,
    )
    doc.outline = _extract_headings(full_text)
    if chunk:
        doc.chunks = _chunk_text(full_text, max_chunk)

    return doc


def _parse_markdown(
    path: str, chunk: bool, max_chunk: int, max_total: int
) -> ParsedDocument:
    """Markdown 文件直接读取。"""
    text = Path(path).read_text(encoding="utf-8")
    doc = ParsedDocument(
        filename=Path(path).name,
        file_type="md",
        total_chars=len(text),
        plain_text=text,
    )
    doc.outline = _extract_headings(text)
    if chunk:
        doc.chunks = _chunk_text(text, max_chunk)
    return doc


def _parse_text(
    path: str, chunk: bool, max_chunk: int, max_total: int
) -> ParsedDocument:
    """纯文本文件。"""
    # 尝试自动检测编码
    text = None
    for enc in ("utf-8", "gbk", "gb2312", "big5", "latin-1"):
        try:
            text = Path(path).read_text(encoding=enc)
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    if text is None:
        raise ValueError(f"无法识别文件编码: {path}")

    doc = ParsedDocument(
        filename=Path(path).name,
        file_type="txt",
        total_chars=len(text),
        plain_text=text,
    )
    doc.outline = _extract_headings(text)
    if chunk:
        doc.chunks = _chunk_text(text, max_chunk)
    return doc


# ---------------------------------------------------------------------------
# 文本工具
# ---------------------------------------------------------------------------

def _chunk_text(text: str, max_chars: int) -> List[DocumentChunk]:
    """将长文本按段落边界切块，尽量避免切断句子。"""
    paragraphs = text.split("\n\n")
    chunks: List[DocumentChunk] = []
    current_chunk: List[str] = []
    current_len = 0
    current_page = 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        para_len = len(para)

        if current_len + para_len > max_chars and current_chunk:
            chunks.append(DocumentChunk(
                index=len(chunks),
                text="\n\n".join(current_chunk),
                page_range=(current_page, current_page),
            ))
            current_chunk = [para]
            current_len = para_len
            current_page += 1
        else:
            current_chunk.append(para)
            current_len += para_len

    if current_chunk:
        chunks.append(DocumentChunk(
            index=len(chunks),
            text="\n\n".join(current_chunk),
            page_range=(current_page, current_page),
        ))

    return chunks


def _extract_headings(text: str) -> List[str]:
    """提取 markdown 风格的标题作为大纲。"""
    headings = []
    for line in text.split("\n"):
        line = line.strip()
        m = re.match(r"^(#{1,4})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            indent = "  " * (level - 1)
            headings.append(f"{indent}- {m.group(2)}")
    return headings


def _truncate_text(text: str, max_chars: int) -> str:
    """在段落边界截断，不切断句子。"""
    if len(text) <= max_chars:
        return text
    trunc = text[:max_chars]
    # 回退到最近的段落结束
    last_para = max(trunc.rfind("\n\n"), trunc.rfind("\n"))
    if last_para > max_chars // 2:
        trunc = text[:last_para]
    return trunc + "\n\n... (内容过长，已截断)"
